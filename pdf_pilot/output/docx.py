"""Word (docx) 输出格式化"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def write_docx(
    md_text: str,
    output_path: str | Path,
) -> Path:
    """将 Markdown 内容转换为 Word 文档

    通过解析 Markdown 基本结构（标题、段落、表格、列表）并生成 docx。

    Args:
        md_text: Markdown 内容
        output_path: 输出文件路径 (.docx)

    Returns:
        Path: 输出文件路径
    """
    from docx import Document
    from docx.shared import Pt

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith("#"):
            level = 0
            for ch in line:
                if ch == "#":
                    level += 1
                else:
                    break
            text = line[level:].strip()
            if text.startswith(" "):
                text = text[1:]
            doc.add_heading(text, level=min(level, 4))

        # 表格行 (检测 Markdown 表格格式)
        elif "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and ("|" in lines[i] or "---" in lines[i] or lines[i].strip() == ""):
                if "|" in lines[i] or "---" in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            _add_table_to_doc(doc, table_lines)
            continue

        # 列表项
        elif line.strip().startswith(("- ", "* ")):
            text = line.strip()[2:].strip()
            doc.add_paragraph(text, style="List Bullet")

        elif line.strip() and line.strip()[0:2] in ("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."):
            text = line.strip().split(". ", 1)[-1] if ". " in line.strip() else line.strip()
            doc.add_paragraph(text, style="List Number")

        # 代码块
        elif line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        # 普通段落
        else:
            paragraph_lines = [line.strip()]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if (
                    not next_line or
                    next_line.startswith("#") or
                    next_line.startswith(("- ", "* ", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")) or
                    next_line.startswith("```")
                ):
                    break
                paragraph_lines.append(next_line)
                i += 1

            doc.add_paragraph(" ".join(paragraph_lines))
            continue

        i += 1

    doc.save(str(output))
    logger.info(f"Word document written to {output}")
    return output


def _add_table_to_doc(doc, table_lines: list[str]):
    """将 Markdown 表格行添加到 Word 文档"""
    # 过滤掉分隔行 (---|---|---)
    data_lines = [
        line for line in table_lines
        if not all(c in ("|", "-", " ", ":") for c in line.strip())
    ]

    if not data_lines:
        return

    # 解析单元格
    def parse_cells(line):
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [cell.strip() for cell in line.split("|")]

    headers = parse_cells(data_lines[0])
    rows = [parse_cells(line) for line in data_lines[1:]]

    # 统一列数
    max_cols = max(len(headers), max([len(r) for r in rows] + [0]))

    # 添加表格
    table = doc.add_table(rows=1 + len(rows), cols=max_cols)
    table.style = "Table Grid"

    # 表头
    for j, cell_text in enumerate(headers):
        table.rows[0].cells[j].text = cell_text[:100]  # 截断过长内容

    # 数据行
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                table.rows[i + 1].cells[j].text = cell_text[:100]
